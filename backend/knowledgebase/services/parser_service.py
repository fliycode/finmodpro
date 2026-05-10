import posixpath
import re
import xml.etree.ElementTree as ET
import zipfile
from io import BytesIO
from pathlib import Path

from django.conf import settings


RELATIONSHIPS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"

_DOCX_STYLE_TO_ELEMENT_MAPPING = {
    "Heading 1": "Title",
    "Heading 2": "Title",
    "Heading 3": "Title",
    "Heading 4": "Title",
    "Heading 5": "Title",
    "Heading 6": "Title",
    "Heading 7": "Title",
    "Heading 8": "Title",
    "Heading 9": "Title",
    "List": "ListItem",
    "List 2": "ListItem",
    "List 3": "ListItem",
    "List Bullet": "ListItem",
    "List Bullet 2": "ListItem",
    "List Bullet 3": "ListItem",
    "List Continue": "ListItem",
    "List Continue 2": "ListItem",
    "List Continue 3": "ListItem",
    "List Number": "ListItem",
    "List Number 2": "ListItem",
    "List Number 3": "ListItem",
    "List Paragraph": "ListItem",
    "Subtitle": "Title",
    "TOCHeading": "Title",
    "Title": "Title",
}

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
_TABLE_ROW_RE = re.compile(r"^\|(.+)\|$")


class ParserService:
    _MULTISPACE_PATTERN = re.compile(r"[^\S\r\n]+")
    _MULTIBLANK_PATTERN = re.compile(r"\n{3,}")

    def clean_text(self, text):
        cleaned = (text or "").replace("\x00", "")
        cleaned = cleaned.replace("\r\n", "\n").replace("\r", "\n")
        cleaned = re.sub(r"(?<=\w)-\n(?=\w)", "", cleaned)
        cleaned = self._MULTISPACE_PATTERN.sub(" ", cleaned)
        cleaned = re.sub(r"[ \t]+\n", "\n", cleaned)
        cleaned = self._MULTIBLANK_PATTERN.sub("\n\n", cleaned)
        return cleaned.strip()

    def _base_document_metadata(self, *, source_parser, source_strategy, fallback_used, element_count=0):
        return {
            "source_parser": source_parser,
            "source_strategy": source_strategy,
            "fallback_used": fallback_used,
            "element_count": element_count,
        }

    def _base_chunk_metadata_defaults(self, *, source_parser, source_strategy):
        return {
            "page_number": None,
            "section_title": "",
            "element_types": [],
            "source_parser": source_parser,
            "source_strategy": source_strategy,
        }

    def _structured_result(
        self,
        *,
        parsed_text,
        source_parser,
        source_strategy,
        fallback_used,
        element_count=0,
        chunk_metadata_defaults=None,
        document_metadata=None,
    ):
        result = {
            "parsed_text": self.clean_text(parsed_text),
            "document_metadata": self._base_document_metadata(
                source_parser=source_parser,
                source_strategy=source_strategy,
                fallback_used=fallback_used,
                element_count=element_count,
            ),
            "chunk_metadata_defaults": self._base_chunk_metadata_defaults(
                source_parser=source_parser,
                source_strategy=source_strategy,
            ),
        }
        if document_metadata:
            result["document_metadata"].update(document_metadata)
            result["document_metadata"]["source_parser"] = source_parser
            result["document_metadata"]["source_strategy"] = source_strategy
            result["document_metadata"]["fallback_used"] = fallback_used
            result["document_metadata"].setdefault("element_count", element_count)
        if chunk_metadata_defaults:
            result["chunk_metadata_defaults"].update(chunk_metadata_defaults)
            result["chunk_metadata_defaults"]["source_parser"] = source_parser
            result["chunk_metadata_defaults"]["source_strategy"] = source_strategy
        return result

    # ── pypdf fallback ────────────────────────────────────────────────────

    def _parse_pdf_via_pypdf(self, file_path):
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise ValueError("当前环境未安装 PDF 解析依赖。") from exc

        reader = PdfReader(str(file_path))
        pages = []
        for page in reader.pages:
            extracted = page.extract_text() or ""
            if extracted.strip():
                pages.append(extracted.strip())
        return "\n\n".join(pages)

    # ── pymupdf4llm PDF parsing ──────────────────────────────────────────

    @staticmethod
    def _markdown_to_elements(md_text, page_number):
        """Convert a Markdown string into element dicts for the chunker."""
        elements = []
        for line in md_text.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue

            heading_match = _HEADING_RE.match(stripped)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                if text:
                    if level == 1:
                        elem_type = "Title"
                    elif level == 2:
                        elem_type = "Section-header"
                    else:
                        elem_type = "Header"
                    elements.append({
                        "type": elem_type,
                        "text": text,
                        "metadata": {"page_number": page_number},
                    })
                continue

            if _TABLE_ROW_RE.match(stripped):
                cells = [c.strip() for c in stripped.strip("|").split("|")]
                cells = [c for c in cells if c and not set(c) <= {"-", ":", " "}]
                if cells:
                    elements.append({
                        "type": "Table",
                        "text": " | ".join(cells),
                        "metadata": {"page_number": page_number},
                    })
                continue

            elements.append({
                "type": "Paragraph",
                "text": stripped,
                "metadata": {"page_number": page_number},
            })

        return elements

    def _parse_pdf_via_pymupdf4llm(self, file_path):
        import pymupdf4llm

        page_chunks = pymupdf4llm.to_markdown(str(file_path), page_chunks=True)
        elements = []
        for chunk in page_chunks:
            md_text = chunk.get("text", "")
            metadata = chunk.get("metadata", {})
            page_number = metadata.get("page", 0) + 1
            elements.extend(self._markdown_to_elements(md_text, page_number))
        return elements

    # ── python-docx DOCX parsing ─────────────────────────────────────────

    @staticmethod
    def _source_part_path(relationships_path):
        if relationships_path == "_rels/.rels":
            return ""
        container, _, relationship_name = relationships_path.partition("/_rels/")
        if not relationship_name.endswith(".rels"):
            return ""
        part_name = relationship_name[:-5]
        if not container:
            return part_name
        return f"{container}/{part_name}"

    @staticmethod
    def _normalize_internal_target(base_dir, target):
        normalized_target = (target or "").strip()
        if not normalized_target or normalized_target.upper() == "NULL":
            return None
        if normalized_target.startswith("/"):
            candidate = normalized_target.lstrip("/")
        else:
            candidate = posixpath.normpath(posixpath.join(base_dir, normalized_target))
        candidate = candidate.lstrip("/")
        if not candidate or candidate == "." or candidate.startswith("../"):
            return None
        return candidate

    @classmethod
    def _repair_docx_relationships(cls, file_bytes):
        source_buffer = BytesIO(file_bytes)
        rewritten_payloads = {}

        with zipfile.ZipFile(source_buffer) as source_zip:
            archive_members = set(source_zip.namelist())
            for info in source_zip.infolist():
                if not info.filename.endswith(".rels"):
                    continue
                payload = source_zip.read(info.filename)
                try:
                    root = ET.fromstring(payload)
                except ET.ParseError:
                    continue
                if root.tag != f"{{{RELATIONSHIPS_NS}}}Relationships":
                    continue

                base_dir = posixpath.dirname(cls._source_part_path(info.filename))
                removed = False

                for relationship in list(root):
                    if relationship.tag != f"{{{RELATIONSHIPS_NS}}}Relationship":
                        continue
                    if relationship.attrib.get("TargetMode") == "External":
                        continue
                    normalized = cls._normalize_internal_target(
                        base_dir,
                        relationship.attrib.get("Target"),
                    )
                    if normalized and normalized in archive_members:
                        continue
                    root.remove(relationship)
                    removed = True

                if removed:
                    rewritten_payloads[info.filename] = ET.tostring(
                        root,
                        encoding="utf-8",
                        xml_declaration=True,
                    )

        if not rewritten_payloads:
            return file_bytes

        source_buffer.seek(0)
        output_buffer = BytesIO()
        with zipfile.ZipFile(source_buffer) as source_zip, zipfile.ZipFile(output_buffer, "w") as target_zip:
            for info in source_zip.infolist():
                payload = rewritten_payloads.get(info.filename, source_zip.read(info.filename))
                target_zip.writestr(info, payload)

        return output_buffer.getvalue()

    @staticmethod
    def _docx_paragraph_element_type(paragraph):
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name in _DOCX_STYLE_TO_ELEMENT_MAPPING:
            return _DOCX_STYLE_TO_ELEMENT_MAPPING[style_name]
        normalized_style = style_name.lower()
        if "heading" in normalized_style or "title" in normalized_style:
            return "Title"
        if normalized_style.startswith("list"):
            return "ListItem"
        return "Paragraph"

    def _parse_docx_via_python_docx(self, file_path):
        from docx import Document

        file_bytes = Path(file_path).read_bytes()
        repaired_bytes = self._repair_docx_relationships(file_bytes)
        doc = Document(BytesIO(repaired_bytes))
        elements = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                elem_type = self._docx_paragraph_element_type(para)
                elements.append({"type": elem_type, "text": text, "metadata": {}})

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    elements.append({
                        "type": "Table",
                        "text": " | ".join(cells),
                        "metadata": {},
                    })

        return elements

    # ── result builders ───────────────────────────────────────────────────

    def _txt_result(self, text):
        return self._structured_result(
            parsed_text=text,
            source_parser="txt",
            source_strategy="local",
            fallback_used=False,
            element_count=0,
        )

    def _pdf_fallback_result(self, text):
        return self._structured_result(
            parsed_text=text,
            source_parser="pypdf",
            source_strategy="fallback",
            fallback_used=True,
            element_count=0,
        )

    def _elements_to_result(self, elements, source_parser, strategy):
        texts = []
        page_numbers = set()
        element_types = {}

        for elem in elements:
            text = (elem.get("text") or "").strip()
            if text:
                texts.append(text)
            meta = elem.get("metadata") or {}
            if meta.get("page_number") is not None:
                page_numbers.add(meta["page_number"])
            elem_type = elem.get("type")
            if elem_type:
                element_types[elem_type] = element_types.get(elem_type, 0) + 1

        parsed_text = "\n\n".join(texts)
        if not parsed_text:
            raise ValueError("解析返回空文本。")

        chunk_defaults = {}
        if page_numbers:
            chunk_defaults["page_number"] = min(page_numbers)
        if element_types:
            chunk_defaults["element_types"] = sorted(element_types.keys())

        result = self._structured_result(
            parsed_text=parsed_text,
            source_parser=source_parser,
            source_strategy=strategy,
            fallback_used=False,
            element_count=len(elements),
            chunk_metadata_defaults=chunk_defaults,
        )
        result["elements"] = elements
        return result

    # ── dispatch ──────────────────────────────────────────────────────────

    def parse(self, document):
        file_path = Path(document.file.path)
        if document.doc_type == "txt":
            return self._txt_result(file_path.read_text(encoding="utf-8"))

        if document.doc_type == "pdf":
            try:
                elements = self._parse_pdf_via_pymupdf4llm(file_path)
                return self._elements_to_result(elements, "pymupdf4llm", "auto")
            except Exception:
                if not settings.PDF_FALLBACK_ENABLED:
                    raise
                return self._pdf_fallback_result(self._parse_pdf_via_pypdf(file_path))

        if document.doc_type == "docx":
            elements = self._parse_docx_via_python_docx(file_path)
            return self._elements_to_result(elements, "python-docx", "native")

        raise ValueError("不支持的文档类型。")


def parse_document_file(document):
    return ParserService().parse(document)
