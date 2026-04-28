import importlib.util
import io
import unittest
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

from docx import Document

APP_PATH = Path(__file__).resolve().parents[1] / "app.py"
MODULE_SPEC = importlib.util.spec_from_file_location("unstructured_light_app", APP_PATH)
APP_MODULE = importlib.util.module_from_spec(MODULE_SPEC)
MODULE_SPEC.loader.exec_module(APP_MODULE)

_extract_docx = APP_MODULE._extract_docx

RELATIONSHIPS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


class ExtractDocxTests(unittest.TestCase):
    def _docx_bytes(self, document):
        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _with_broken_header_relationship(self, docx_bytes):
        input_buffer = io.BytesIO(docx_bytes)
        output_buffer = io.BytesIO()
        mutated = False

        with zipfile.ZipFile(input_buffer) as source_zip, zipfile.ZipFile(output_buffer, "w") as target_zip:
            for info in source_zip.infolist():
                payload = source_zip.read(info.filename)
                if info.filename == "word/_rels/document.xml.rels":
                    root = ET.fromstring(payload)
                    for relationship in root.findall(f"{{{RELATIONSHIPS_NS}}}Relationship"):
                        if relationship.attrib.get("Type", "").endswith("/header"):
                            relationship.set("Target", "NULL")
                            mutated = True
                            break
                    payload = ET.tostring(root, encoding="utf-8", xml_declaration=True)

                target_zip.writestr(info, payload)

        self.assertTrue(mutated, "expected a header relationship to mutate in the test fixture")
        return output_buffer.getvalue()

    def test_extract_docx_ignores_missing_internal_relationship_targets(self):
        document = Document()
        document.sections[0].header.paragraphs[0].text = "Internal only"
        document.add_heading("Quarterly Risk Review", level=1)
        document.add_paragraph("Liquidity stayed stable.")

        broken_docx = self._with_broken_header_relationship(self._docx_bytes(document))

        elements = _extract_docx(broken_docx)

        self.assertEqual(
            [(element["type"], element["text"]) for element in elements],
            [
                ("Title", "Quarterly Risk Review"),
                ("Paragraph", "Liquidity stayed stable."),
            ],
        )

    def test_extract_docx_maps_list_styles_to_list_items(self):
        document = Document()
        document.add_paragraph("First risk", style="List Bullet")
        document.add_paragraph("Second risk", style="List Number")

        elements = _extract_docx(self._docx_bytes(document))

        self.assertEqual([element["type"] for element in elements], ["ListItem", "ListItem"])


if __name__ == "__main__":
    unittest.main()
