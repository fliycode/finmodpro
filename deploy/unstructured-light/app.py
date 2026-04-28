"""Lightweight document text extraction — strategy=fast only.

Covers ``pdf`` (via pdfminer.six), ``docx`` (via python-docx), and
``txt`` (raw read).  No ``unstructured``, Detectron2, PyTorch, Tesseract,
or LibreOffice — total image ~200 MB.
"""

import posixpath
import xml.etree.ElementTree as ET
import zipfile
from io import BytesIO

from fastapi import FastAPI, File, HTTPException, Query, UploadFile

app = FastAPI()

RELATIONSHIPS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
STYLE_TO_ELEMENT_MAPPING = {
    "Heading 1": "Title",
    "Heading 2": "Title",
    "Heading 3": "Title",
    "Heading 4": "Title",
    "Heading 5": "Title",
    "Heading 6": "Title",
    "Heading 7": "Title",
    "Heading 8": "Title",
    "Heading 9": "Title",
    "List": "ListItem",
    "List 2": "ListItem",
    "List 3": "ListItem",
    "List Bullet": "ListItem",
    "List Bullet 2": "ListItem",
    "List Bullet 3": "ListItem",
    "List Continue": "ListItem",
    "List Continue 2": "ListItem",
    "List Continue 3": "ListItem",
    "List Number": "ListItem",
    "List Number 2": "ListItem",
    "List Number 3": "ListItem",
    "List Paragraph": "ListItem",
    "Subtitle": "Title",
    "TOCHeading": "Title",
    "Title": "Title",
}

# ── helpers ──────────────────────────────────────────────────────────────


def _extract_pdf(file_bytes: bytes) -> list[dict]:
    from pdfminer.high_level import extract_text  # noqa: E402

    text = extract_text(BytesIO(file_bytes))
    elements = []
    for line in text.splitlines():
        stripped = line.strip()
        if stripped:
            elements.append({"type": "Paragraph", "text": stripped, "metadata": {}})
    if not elements:
        # pdfminer returned nothing — basic fallback.
        decoded = text.strip()
        if decoded:
            elements.append({"type": "Paragraph", "text": decoded, "metadata": {}})
    return elements


def _source_part_path(relationships_path: str) -> str:
    if relationships_path == "_rels/.rels":
        return ""

    container, _, relationship_name = relationships_path.partition("/_rels/")
    if not relationship_name.endswith(".rels"):
        return ""

    part_name = relationship_name[:-5]
    if not container:
        return part_name
    return f"{container}/{part_name}"


def _normalize_internal_target(base_dir: str, target: str | None) -> str | None:
    normalized_target = (target or "").strip()
    if not normalized_target or normalized_target.upper() == "NULL":
        return None

    if normalized_target.startswith("/"):
        candidate = normalized_target.lstrip("/")
    else:
        candidate = posixpath.normpath(posixpath.join(base_dir, normalized_target))

    candidate = candidate.lstrip("/")
    if not candidate or candidate == "." or candidate.startswith("../"):
        return None
    return candidate


def _repair_docx_relationships(file_bytes: bytes) -> bytes:
    source_buffer = BytesIO(file_bytes)
    rewritten_payloads: dict[str, bytes] = {}

    with zipfile.ZipFile(source_buffer) as source_zip:
        archive_members = set(source_zip.namelist())
        for info in source_zip.infolist():
            if not info.filename.endswith(".rels"):
                continue

            payload = source_zip.read(info.filename)
            try:
                root = ET.fromstring(payload)
            except ET.ParseError:
                continue

            if root.tag != f"{{{RELATIONSHIPS_NS}}}Relationships":
                continue

            base_dir = posixpath.dirname(_source_part_path(info.filename))
            removed_invalid_relationship = False

            for relationship in list(root):
                if relationship.tag != f"{{{RELATIONSHIPS_NS}}}Relationship":
                    continue
                if relationship.attrib.get("TargetMode") == "External":
                    continue

                normalized_target = _normalize_internal_target(
                    base_dir,
                    relationship.attrib.get("Target"),
                )
                if normalized_target and normalized_target in archive_members:
                    continue

                root.remove(relationship)
                removed_invalid_relationship = True

            if removed_invalid_relationship:
                rewritten_payloads[info.filename] = ET.tostring(
                    root,
                    encoding="utf-8",
                    xml_declaration=True,
                )

    if not rewritten_payloads:
        return file_bytes

    source_buffer.seek(0)
    output_buffer = BytesIO()
    with zipfile.ZipFile(source_buffer) as source_zip, zipfile.ZipFile(output_buffer, "w") as target_zip:
        for info in source_zip.infolist():
            payload = rewritten_payloads.get(info.filename, source_zip.read(info.filename))
            target_zip.writestr(info, payload)

    return output_buffer.getvalue()


def _paragraph_element_type(paragraph) -> str:
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name in STYLE_TO_ELEMENT_MAPPING:
        return STYLE_TO_ELEMENT_MAPPING[style_name]

    normalized_style = style_name.lower()
    if "heading" in normalized_style or "title" in normalized_style:
        return "Title"
    if normalized_style.startswith("list"):
        return "ListItem"
    return "Paragraph"


def _extract_docx(file_bytes: bytes) -> list[dict]:
    from docx import Document  # noqa: E402

    repaired_bytes = _repair_docx_relationships(file_bytes)
    doc = Document(BytesIO(repaired_bytes))
    elements = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            elem_type = _paragraph_element_type(para)
            elements.append({"type": elem_type, "text": text, "metadata": {}})

    # Also extract tables.
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                elements.append(
                    {"type": "Table", "text": " | ".join(cells), "metadata": {}}
                )
    return elements


# ── API ──────────────────────────────────────────────────────────────────


@app.post("/general/v0/general")
def parse_file(
    files: UploadFile = File(...),
    strategy: str = Query("fast"),
):
    if strategy not in ("fast", "auto"):
        raise HTTPException(
            status_code=400,
            detail=f"strategy={strategy} not supported (only fast/auto).",
        )

    content_type = files.content_type or ""
    filename = (files.filename or "").lower()

    raw = files.file.read()

    try:
        if "pdf" in content_type or filename.endswith(".pdf"):
            elements = _extract_pdf(raw)
        elif filename.endswith(".docx") or "wordprocessingml" in content_type:
            elements = _extract_docx(raw)
        elif filename.endswith(".txt") or "text/plain" in content_type:
            text = raw.decode("utf-8", errors="replace")
            elements = [
                {"type": "Paragraph", "text": line.strip(), "metadata": {}}
                for line in text.splitlines()
                if line.strip()
            ]
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported type: {content_type}")

        return elements
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.get("/healthcheck")
def healthcheck():
    return {"status": "ok"}
